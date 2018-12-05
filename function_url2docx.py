def url2docx(url,filename):
    #this function is for wyckofftrade.com blog page download,if other website,need make some modify
    import requests
    from bs4 import BeautifulSoup
    from docx import Document
    from docx.shared import Inches
    document = Document()
    sections = document.sections
    section = sections[0]
    pagewidth = section.page_width
    left_margin = section.left_margin
    section.left_margin = left_margin/4
    section.right_margin = left_margin/4
    picturewidth = pagewidth-left_margin/2
    res = requests.get(url)
    res.encoding=res.apparent_encoding
    
    soup = BeautifulSoup(res.text,'html.parser')
    c1 = soup.find('div',class_='post-content').children
    for i in c1:
        if i.name=='h1' or i.name=='h2':
            document.add_heading(i.string)
        if i.name=='p' :
            #print(i.string)
            document.add_paragraph(i.string)
            for j in  i.children:
                if j.name=='img':
                    #print(j.get('src'))
                    res = requests.get(j.get('src'))
                    if res.status_code==200:
                        open('hahaha.jpg','wb').write(res.content)
                        document.add_picture('hahaha.jpg',width=picturewidth)
    filename = filename+'.docx'
    document.save(filename)     