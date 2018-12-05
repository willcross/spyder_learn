import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
document = Document()
sections = document.sections
section = sections[0]
pagewidth = section.page_width
left_margin = section.left_margin
section.left_margin = left_margin/4
section.right_margin = left_margin/4
picturewidth = pagewidth-left_margin/2
res = requests.get('https://wyckofftrade.com/blog/powerwave.html')
res.encoding=res.apparent_encoding
#soup = BeautifulSoup(res.text,'lxml')
soup = BeautifulSoup(res.text,'html.parser')
#soup = BeautifulSoup(res.content,'html.parser')
c1 = soup.find('div',class_='post-content').children
# for i in c1:
    # try:
        # #print(i.text)
        # print(i.string)
        
    # except:
        # pass


c1 = soup.find('div',class_='post-content').children
for i in c1:
    if i.name=='p' or i.name=='h1' or i.name=='h2':
        #print(i.string)
        document.add_paragraph(i.string)
        for j in  i.children:
            if j.name=='img':
                #print(j.get('src'))
                res = requests.get(j.get('src'))
                if res.status_code==200:
                    open('hahaha.jpg','wb').write(res.content)
                    document.add_picture('hahaha.jpg',width=picturewidth)
document.save('powerwave2.docx')         